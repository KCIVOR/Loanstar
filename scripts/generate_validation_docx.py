"""Generate LoanStar MVP Validation Guide as DOCX."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "E8E8E8")
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = val
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def main() -> None:
    out_path = Path(__file__).resolve().parents[1] / "docs" / "LoanStar_MVP_Validation_Guide.docx"
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("LoanStar LMS — MVP Validation Guide", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "End-to-end user journey for validating the happy path (P1–P8). "
        "Version 1.0 — July 2026"
    )
    doc.add_paragraph()

    doc.add_heading("Overview", level=1)
    doc.add_paragraph(
        "The MVP is built through Phase 8 (P1–P8). Automated unit tests should pass 25/25 "
        "before manual validation. This guide provides phase-by-phase checklists and a full "
        "sample run using the F2 Del Poso computation fixture."
    )

    doc.add_heading("Before You Start (15 minutes)", level=1)

    doc.add_heading("1. Automated baseline", level=2)
    doc.add_paragraph("From the loanstar/ folder, run:")
    p = doc.add_paragraph()
    run = p.add_run("npm test")
    run.bold = True
    doc.add_paragraph("Pass if: 25 tests, 0 failures (computation F1–F7, BLRI F2, RBAC, triggers, TAT).")

    doc.add_heading("2. Dev server", level=2)
    doc.add_paragraph("App URL: http://localhost:3000 (npm run dev in loanstar/).")

    doc.add_heading("3. Create test staff accounts", level=2)
    doc.add_paragraph(
        "Log in as Super Admin → Admin → Users → create one user per role. "
        "Use separate browser profiles or incognito windows per role."
    )
    add_table(
        doc,
        ["Email (suggested)", "Role(s)", "Portal"],
        [
            ["admin@loanstar.local", "Super Admin", "/admin"],
            ["csa@loanstar.test", "CSA", "/csa"],
            ["cig@loanstar.test", "CIG", "/cig"],
            ["committee1@loanstar.test", "Committee", "/committee"],
            ["committee2@loanstar.test", "Committee (optional)", "/committee"],
            ["committee3@loanstar.test", "Committee (optional)", "/committee"],
            ["lra@loanstar.test", "LRA", "/lra"],
            ["ar@loanstar.test", "AR", "/ar"],
            ["collector@loanstar.test", "Collector", "/collector"],
            ["remedial@loanstar.test", "Remedial", "/remedial"],
            ["agent@loanstar.test", "Agent", "/agent"],
            ["rvckmlnrmsnt@gmail.com", "Borrower (existing)", "/borrower"],
        ],
    )
    doc.add_paragraph("Suggested test password: LoanStar2026!")

    doc.add_heading("4. Add F2 validation loan type", level=2)
    doc.add_paragraph("Admin → Loan Types → add:")
    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Name", "F2-VALIDATION"],
            ["Interest rate", "2.10% (0.021)"],
            ["PF rate", "11.34% (0.1134)"],
            ["Effective from", "Today"],
            ["Active", "Yes"],
        ],
    )

    doc.add_heading("5. Prepare dummy files", level=2)
    doc.add_paragraph("Keep 10+ small PDFs or images ready for document uploads.")

    doc.add_heading("Lifecycle Map", level=1)
    add_bullets(
        doc,
        [
            "Borrower registers → Upload + sign intake docs",
            "CSA computes + NCL + endorse → CIG verifies + auto-forward",
            "Committee votes + final Approve → CSA discloses → Borrower signs computation",
            "LRA path + PDC + generate docs → Borrower signs release docs + briefing",
            "LRA releases + closes → AR masterlist + assign collector",
            "Borrower pays → Collector DCR → AR posts → Loan Active + Reports",
        ],
    )

    doc.add_heading("Phase-by-Phase Validation Checklist", level=1)

    doc.add_heading("P1 — Foundation (admin@loanstar.local)", level=2)
    add_table(
        doc,
        ["#", "Step", "Pass if"],
        [
            ["1", "Open /admin", "Dashboard loads, no errors"],
            ["2", "Roles → open CSA role", "Module permissions visible"],
            ["3", "Users → create csa@loanstar.test", "User saved with CSA role"],
            ["4", "Config", "Penalty 5%, coverage 35%, aging thresholds visible"],
            ["5", "Loan Types", "22 active types; inactive not selectable in CSA"],
            ["6", "Audit", "Page loads; events appear after workflow actions"],
        ],
    )

    doc.add_heading("P2 — Borrower & Agent", level=2)
    add_table(
        doc,
        ["#", "Step", "Pass if"],
        [
            ["7", "/borrower/register or use existing BN300001", "Borrower number assigned"],
            ["8", "Borrower uploads intake docs", "Status badges move to Uploaded"],
            ["9", "Borrower clicks Sign on each doc", "All required items Confirmed"],
            ["10", "Borrower status timeline", "Documents Pending → Submitted"],
            ["11", "Agent /agent → create lead", "Checklist flags only; no file content"],
        ],
    )

    doc.add_heading("P3 — CSA Intake + Computation", level=2)
    add_table(
        doc,
        ["#", "Step", "Pass if"],
        [
            ["12", "CSA /csa → open application", "Profile + checklist visible"],
            ["13", "Record NCL pass", "NCL shows pass"],
            ["14", "Compute F2 sample (see below)", "Principal 102,605.05; monthly 17,428.20"],
            ["15", "Borrower signs computation", "signedAt populated"],
            ["16", "Endorse to CIG", "Status → For Verification; gate enforced"],
            ["17", "CSA tries to edit after endorse", "Fields locked / 403"],
        ],
    )

    doc.add_heading("P4 — CIG Verification", level=2)
    add_table(
        doc,
        ["#", "Step", "Pass if"],
        [
            ["18", "CIG /cig → file in queue", "Endorsed application visible"],
            ["19", "Fill form + record 5 checks", "NFIS, MF, LSLG, POEA, Marina saved"],
            ["20", "Save complete form", "Auto-forwards to Committee"],
            ["21", "Status", "For Approval"],
            ["22", "Borrower view", "Status label only — no verification notes"],
        ],
    )

    doc.add_heading("P5 — Committee + Negotiation", level=2)
    add_table(
        doc,
        ["#", "Step", "Pass if"],
        [
            ["23", "Committee votes (1–3 users)", "Tally shows majority; file does NOT move yet"],
            ["24", "Final action → Approve", "Status → Approved"],
            ["25", "CSA Disclose approved terms", "Status → Awaiting Confirmation"],
            ["26", "Borrower signs computation", "Status → LRA Pending; file in LRA queue"],
            ["27", "Borrower view", "No votes, tallies, or committee comments"],
        ],
    )

    doc.add_heading("P6 — LRA Release", level=2)
    add_table(
        doc,
        ["#", "Step", "Pass if"],
        [
            ["28", "LRA Start processing", "Release file created"],
            ["29", "Choose With PDC → encode PDC", "PDC rows saved"],
            ["30", "Generate release documents", "BLRI, PN, Disclosure, etc. listed"],
            ["31", "Borrower signs each generated doc", "All show Signed"],
            ["32", "Borrower signs briefing", "Status → Release — Ready for Disbursement"],
            ["33", "LRA Record release", "Status → Released"],
            ["34", "Upload Signed Check Voucher → Close file", "Status → Closed — Transmitted"],
        ],
    )

    doc.add_heading("P7 — AR + Collection", level=2)
    add_table(
        doc,
        ["#", "Step", "Pass if"],
        [
            ["35", "AR /ar", "Masterlist row auto-created"],
            ["36", "Assign Portfolio A + collector", "Collector sees account only on their dashboard"],
            ["37", "Borrower uploads payment proof", "Status Pending verification"],
            ["38", "Collector confirms → DCR → Submit", "DCR in AR queue"],
            ["39", "AR /ar/dcr → reconcile", "Payment Posted/Paid; balance updates"],
            ["40", "Borrower Loan Active panel", "7 × 17,428.20 schedule visible"],
        ],
    )

    doc.add_heading("P8 — Reports + QA", level=2)
    add_table(
        doc,
        ["#", "Step", "Pass if"],
        [
            ["41", "/reports (AR or Super Admin)", "Pipeline, aging, TAT load"],
            ["42", "Admin → Audit", "Trigger events logged"],
            ["43", "Confidentiality spot-checks", "Boundaries hold (see below)"],
        ],
    )

    doc.add_page_break()
    doc.add_heading("End-to-End Sample Run — F2 Del Poso", level=1)
    doc.add_paragraph(
        "Sample borrower: Del Poso seafarer loan. You may continue with existing borrower "
        "Rovick Romasanta (BN300001, status Documents Pending) or register fresh."
    )

    doc.add_heading("Sample Data Card", level=2)
    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Loan type", "F2-VALIDATION"],
            ["Input mode", "NET_SARADO"],
            ["Amount", "90,000.00"],
            ["Terms", "7"],
            ["Addon months", "2"],
            ["Release path", "With PDC"],
            ["PDC monthly amount", "17,428.20"],
            ["First check date", "2026-08-10 (or any future date)"],
            ["Bank", "BDO (any)"],
            ["Payment proof ref", "DEP-2026-001"],
        ],
    )

    doc.add_heading("Expected Computation (verify to the centavo)", level=2)
    add_table(
        doc,
        ["Line", "Amount"],
        [
            ["Principal", "102,605.05"],
            ["Admin Cost", "3,421.90"],
            ["Total interest", "19,392.36"],
            ["Total loan", "121,997.41"],
            ["Monthly amortization", "17,428.20"],
            ["Net released", "90,000.00"],
        ],
    )

    doc.add_heading("Act 1 — Borrower (15–20 min)", level=2)
    doc.add_paragraph("Login: rvckmlnrmsnt@gmail.com (or new borrower)")
    add_numbered(
        doc,
        [
            "Go to http://localhost:3000/borrower",
            "Open the application (BN300001)",
            "For each required intake document: Upload → pick dummy PDF/image",
            "Click Sign → confirm signature on each uploaded doc",
            "Pass: All required items Confirmed; status moves toward Submitted",
        ],
    )

    doc.add_heading("Act 2 — CSA (10–15 min)", level=2)
    doc.add_paragraph("Login: csa@loanstar.test → /csa → open application")
    add_numbered(
        doc,
        [
            "Review/save borrower profile",
            "Record NCL pass",
            "Computation: F2-VALIDATION, NET_SARADO, 90000, Terms 7, Addon 2 → Compute",
            "Verify amounts match expected computation table",
            "Borrower signs computation (switch login)",
            "Back as CSA → Endorse to CIG",
            "Pass: Status = For Verification",
        ],
    )

    doc.add_heading("Act 3 — CIG (10 min)", level=2)
    doc.add_paragraph("Login: cig@loanstar.test → /cig")
    add_numbered(
        doc,
        [
            "Open the application",
            "Complete verification sections (field completeness, PIC, crewing manager, refs)",
            "Set finding Positive or Negative (both should forward)",
            "Record each check: NFIS, MF, LSLG, POEA, Marina → Pass",
            "Save form",
            "Pass: Status = For Approval",
        ],
    )

    doc.add_heading("Act 4 — Committee (5–10 min)", level=2)
    doc.add_paragraph("Login: committee1@loanstar.test → /committee")
    add_numbered(
        doc,
        [
            "Open application → review computation + CIG findings",
            "Cast vote: Approve",
            "(Optional) Log in as committee2/3 — confirm tally shows majority but status unchanged",
            "Click Final action → Approve",
            "Pass: Status = Approved; TAT days visible",
        ],
    )

    doc.add_heading("Act 5 — Negotiation (5 min)", level=2)
    add_numbered(
        doc,
        [
            "CSA: Disclose approved terms",
            "Borrower: Sign computation in negotiation panel",
            "Pass: Status = LRA Pending; file in /lra queue",
        ],
    )

    doc.add_heading("Act 6 — LRA Release (20–25 min)", level=2)
    doc.add_paragraph("Login: lra@loanstar.test → /lra → open application")
    add_numbered(
        doc,
        [
            "Start processing",
            "Release path: With PDC → Save",
            "PDC encoding: amount 17428.20, bank, check dates/numbers → Save",
            "Generate release documents",
            "Borrower: sign every generated document + BriefingSign",
            "LRA: Record release (enabled only after briefing)",
            "Upload Signed Check Voucher on release checklist → Close file",
            "Pass: Status = Closed — Transmitted; BLRI shows F2 amounts",
        ],
    )

    doc.add_heading("Act 7 — AR + Collection (15 min)", level=2)
    doc.add_paragraph("AR (ar@loanstar.test):")
    add_numbered(
        doc,
        [
            "/ar → confirm masterlist row exists",
            "Assign Portfolio A + collector@loanstar.test",
        ],
    )
    doc.add_paragraph("Borrower:")
    add_numbered(
        doc,
        [
            "Open application → Loan Active panel",
            "Upload payment proof: 17,428.20, ref DEP-2026-001",
        ],
    )
    doc.add_paragraph("Collector (collector@loanstar.test):")
    add_numbered(
        doc,
        [
            "/collector → see assigned account only",
            "Confirm pending payment",
            "Start DCR → select payment → Submit DCR",
        ],
    )
    doc.add_paragraph("AR:")
    add_numbered(
        doc,
        [
            "/ar/dcr → deposit ref DEP-2026-001 → Reconcile/Post",
            "Pass: Payment Posted/Paid; schedule 7 × 17,428.20",
        ],
    )

    doc.add_heading("Act 8 — Reports + Audit (5 min)", level=2)
    doc.add_paragraph("Login: admin@loanstar.local or ar@loanstar.test")
    add_numbered(
        doc,
        [
            "/reports — pipeline, aging, TAT sections load",
            "/admin/audit — endorse, committee, release, DCR post events visible",
        ],
    )

    doc.add_page_break()
    doc.add_heading("Confidentiality Spot-Checks", level=1)
    add_table(
        doc,
        ["Test", "How", "Pass if"],
        [
            ["Agent boundary", "Agent views lead checklist", "Flags only; cannot open document bytes"],
            ["Borrower boundary", "Borrower page for committee data", "No votes or CIG notes"],
            ["CIG computation", "CIG tries to edit computation", "Read-only / rejected"],
            ["Collector RLS", "Unassigned collector logs in", "Does not see other accounts"],
        ],
    )

    doc.add_heading("Pass/Fail Log", level=1)
    add_bullets(
        doc,
        [
            "[ ] npm test — 25/25 green",
            "[ ] All test users created",
            "[ ] P2 intake docs confirmed",
            "[ ] P3 F2 computation exact match",
            "[ ] P3 endorse gate works",
            "[ ] P4 CIG auto-forward",
            "[ ] P5 committee final action only advances file",
            "[ ] P6 briefing gate blocks release until signed",
            "[ ] P6 close → AR queue",
            "[ ] P7 borrower upload alone does NOT set Paid",
            "[ ] P7 AR reconcile sets Paid",
            "[ ] P8 reports + audit populated",
            "[ ] No console errors on critical pages",
        ],
    )

    doc.add_heading("Known MVP Limits (Not Validation Failures)", level=1)
    add_bullets(
        doc,
        [
            "No real NFIS/bank/POEA integrations (manual pass/fail + proof upload)",
            "Some LRA docs are upload/confirm until PDF templates are provided",
            "Full Playwright E2E needs staging credentials (E2E_BORROWER_EMAIL)",
            "91-day remedial turnover needs date simulation or DB aging override",
            "Committee tally demo needs 3 separate committee logins",
        ],
    )

    doc.add_heading("Recommended Order Today", level=1)
    add_numbered(
        doc,
        [
            "Finish Act 1 (upload + sign intake docs) on existing borrower BN300001",
            "Run Acts 2–8 in sequence with test staff accounts",
            "Fill in the pass/fail log as you go",
        ],
    )

    doc.save(out_path)
    print(f"Created: {out_path}")


if __name__ == "__main__":
    main()
